import os
import asyncio
import datetime
from fastapi import APIRouter, HTTPException
from langchain.text_splitter import RecursiveCharacterTextSplitter
from docx import Document
from pydantic import BaseModel, HttpUrl, Field
from dotenv import load_dotenv
from best_first import BestFirstCrawl
from depth_first import DepthFirstCrawl

load_dotenv()

router = APIRouter()

asyncio.set_event_loop_policy(asyncio.WindowsProactorEventLoopPolicy())


class CrawlRequest(BaseModel):
    url: str
    strategy:str
    method: str
    depth: int = Field(..., ge=0, le=3)

def save_results_to_docx(strategy, method, results: list[dict]) -> str:
    local_storage_path = os.path.join(os.path.expanduser("~"), "Downloads", "crawl_exports")
    os.makedirs(local_storage_path, exist_ok=True)
    timestamp = datetime.datetime.now().strftime("%H%M%S")
    file_path = os.path.join(local_storage_path, f"crawl_result_{strategy}_{method}_{timestamp}.docx")

    doc = Document()
    doc.add_heading("Crawled Content", level=1)

    for i, item in enumerate(results, start=1):
        doc.add_heading(f"{i}. {item['url']}", level=2)
        
        # Use markdown instead of missing 'text'
        text = item.get("fit_markdown") 
        if text:
            text_splitter = RecursiveCharacterTextSplitter(chunk_size=5000, chunk_overlap=1000, separators=["\n\n", "\n", " ", "", "."])
            chunks = text_splitter.split_text(text)
            for chunk in chunks:
                doc.add_paragraph(chunk)
        else:
            doc.add_paragraph("No content available.")

    print(f"Saving DOCX to: {file_path}")
    doc.save(file_path)
    return file_path

@router.post("/")
def start_crawling(request: CrawlRequest):
    try:
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)

        url = request.url
        method = request.method
        strategy = request.strategy
        depth = request.depth
        results = []
        if strategy == "best first":

            if method == "single":

                crawl_single_page_service = BestFirstCrawl()
                crawl_single_page = crawl_single_page_service.crawl_single_page(url)
                results = loop.run_until_complete(crawl_single_page)
               
            elif method == "recursive":

                # Use the best-first crawling strategy
                best_first_crawl_service = BestFirstCrawl()
                best_first_crawl = best_first_crawl_service.best_first_crawl(url,depth)
                results = loop.run_until_complete(best_first_crawl)

                
        elif strategy == "depth first":

            if method == "single":

                crawl_single_page_service = DepthFirstCrawl()
                crawl_single_page = crawl_single_page_service.crawl_single_page(url)
                results = loop.run_until_complete(crawl_single_page)
                
            if method == "recursive":
                #Use the depth-first crawling strategy

                depth_first_crawl_service = DepthFirstCrawl()
                depth_first_crawl = depth_first_crawl_service.depth_first_crawl(url,depth)
                results = loop.run_until_complete(depth_first_crawl)

                
        else:

            raise HTTPException(status_code=400, detail="Invalid strategy")
            
        save_results_to_docx(strategy,method,results)
        return {
            "message": "Crawling completed and data stored successfully.",
            "strategy":strategy,
            "method":method,
            "pages_crawled": len(results or [])
            
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

